# """
# Document loader for the RAG knowledge base.

# This module reads medication guideline files in different formats,
# normalizes their text, and splits them into searchable chunks.
# """

# #import necessary libraries and packages
# from dataclasses import dataclass
# from pathlib import Path
# import json
# import re
# import docx
# import pdfplumber


# @dataclass
# class DocumentChunk:
#     """
#     A single, searchable chunk of text from a source document.
#     """
#     chunk_id:str
#     source_file:str
#     page: int | None
#     text: str
#     metadata: dict

# # To read the files in txt format in the files subfolder
# def read_txt_file(path: Path) -> list[tuple[int, str]]:
#     """Read a TXT file and return a single document with page=None."""
#     text = path.read_text(encoding="utf-8", errors="ignore")
#     return [(None, text)]


# # To read the documents in pdf format
# def read_pdf_file(path: Path) -> list[tuple[int, str]]:
#     """Read each page from a PDF and return a page-indexed list."""
#     pages = []
#     with pdfplumber.open(path) as pdf:
#         for page_number, page in enumerate(pdf.pages, start=1):
#             page_text = page.extract_text() or ""
#             if page_text.strip():
#                 pages.append((page_number, page_text))
#                 return pages
            

# # To read documents in pdf format
# def read_docx_file(path: Path) -> list[tuple[int, str]]:
#     """Read a DOCX file and return one combined page-like document."""
#     document = docx.Document(path)
#     paragraphs = [para.txt for para in document.paragraphs if para.text.strip()]
#     text = "\n".join(paragraphs)
#     return [(None, text)]

# # To read documents in json format
# def read_json_file(path: Path) -> list[tuple[int, str]]:
#     """Read JSON content and convert it to a normalized text representation."""
#     raw = json.loads(path.read_text(encoding="utf-8"))
#     if isinstance(raw, dict):
#         text = json.dumps(raw, indent=2)
#     elif isinstance(raw, list):
#         text = "\n\n".join(json.dumps(item, indent=2) for item in raw)
#     else:
#         text = str(raw)
#         return [(None, text)]
    
# # To put all texts in standard format
# def normalize_text(text: str) -> str:
#     """Remove extra whitespace and normalize newlines for searching."""
#     text = text.replace("\r\n", "\n").replace("\r", "\n")
#     text = re.sub(r"\n{2,}", "\n\n", text)
#     text = re.sub(r"[ \t]+", " ", text)
#     return text.strip

# # Function to chunk documents
# def chunk_text(text: str, source_file, str, page: int | None = None, chunk_size = 350, chunk_overlap = 80,) -> list[DocumentChunk]:
#     """
#     Split a long text string into smaller chunk objects.

#     This helps retrieval when documents are large.
#     """
#     words = text.split()
#     chunks = []
#     start = 0
#     chunk_index = 0

#     while start < len(words):
#         end = min(start + chunk_size, len(words))
#         chunk_words = words[start:end]
#         chunk_text = " ".join(chunk_words).strip()
#         if chunk_text:
#             chunks.append(DocumentChunk(
#                 chunk_id = f"{source_file}:{page or 0}:{chunk_index}",
#                 source_file=source_file,
#                 page = page,
#                 text = chunk_text,
#                 metadata={"source_file": source_file, "page": page, "chunk_index":chunk_index,},
#             ))
#             chunk_index += 1
#             start += chunk_size - chunk_overlap
#             return chunks

# # Function to load the files
# def load_documents(directory: str | Path) -> list[DocumentChunk]:
#     """
#     Load all supported files from the specified directory and return chunks.
#     """
#     base_path = Path(directory)
#     chunks = []

#     for path in sorted(base_path.iterdir()):
#         if not path.is_file():
#             continue

#         extension = path.suffix.lower()
#         if extension == ".txt":
#             pages = read_txt_file(path)
#         elif extension == ".pdf":
#             pages = read_pdf_file(path)
#         elif extension == ".docx":
#             pages = read_docx_file(path)
#         elif extension == ".json":
#             pages = read_json_file(path)
#         else:
#             continue

#         for page_number, raw_text in pages:
#             normalized = normalize_text(raw_text)
#             if not normalized:
#                 continue
#             chunks.extend(chunk_text(normalized, source_file=path.name, page=page_number))
#     return chunks
        
        


"""
Load medication guidance documents from multiple file formats.

This supports TXT, PDF, DOCX, and JSON files.
"""

from __future__ import annotations

import json
from pathlib import Path
from typing import List, Optional, Tuple

import docx
import pdfplumber

from .models import DocumentRecord


def _load_txt(path: Path) -> List[Tuple[Optional[int], str]]:
    """Read plain text files and return a single page-like tuple."""
    content = path.read_text(encoding="utf-8", errors="ignore")
    return [(None, content)]


def _load_pdf(path: Path) -> List[Tuple[Optional[int], str]]:
    """Extract text from each page of a PDF file."""
    pages: List[Tuple[Optional[int], str]] = []
    with pdfplumber.open(path) as pdf:
        for page_number, page in enumerate(pdf.pages, start=1):
            text = page.extract_text() or ""
            if text.strip():
                pages.append((page_number, text))
    return pages


def _load_docx(path: Path) -> List[Tuple[Optional[int], str]]:
    """Extract all paragraphs from a DOCX file into one text blob."""
    document = docx.Document(path)
    paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    text = "\n".join(paragraphs)
    return [(None, text)]


def _load_json(path: Path) -> List[Tuple[Optional[int], str]]:
    """Convert JSON content into a normalized text representation."""
    raw = json.loads(path.read_text(encoding="utf-8"))
    if isinstance(raw, dict):
        text = json.dumps(raw, indent=2)
    elif isinstance(raw, list):
        text = "\n\n".join(json.dumps(item, indent=2) for item in raw)
    else:
        text = str(raw)
    return [(None, text)]


def _normalize_text(text: str) -> str:
    """
    Normalize whitespace and blank lines.

    This makes the source more consistent for chunking and retrieval.
    """
    normalized = text.replace("\r\n", "\n").replace("\r", "\n")
    normalized = "\n".join(line.strip() for line in normalized.splitlines())
    normalized = "\n\n".join(block for block in normalized.split("\n\n") if block.strip())
    return normalized.strip()


def load_documents(directory: Path | str) -> List[DocumentRecord]:
    """
    Load supported documents from the given directory.

    Returns a list of DocumentRecord objects.
    """
    base_path = Path(directory)
    records: List[DocumentRecord] = []

    for path in sorted(base_path.iterdir()):
        if not path.is_file():
            continue

        suffix = path.suffix.lower()
        if suffix == ".txt":
            pages = _load_txt(path)
        elif suffix == ".pdf":
            pages = _load_pdf(path)
        elif suffix == ".docx":
            pages = _load_docx(path)
        elif suffix == ".json":
            pages = _load_json(path)
        else:
            # Skip unsupported file formats
            continue

        for page_number, raw_text in pages:
            normalized = _normalize_text(raw_text)
            if not normalized:
                continue

            records.append(
                DocumentRecord(
                    doc_id=f"{path.name}:{page_number or 0}",
                    source_file=path.name,
                    page=page_number,
                    text=normalized,
                    metadata={
                        "source_file": path.name,
                        "page": page_number,
                        "path": str(path),
                    },
                )
            )

    return records